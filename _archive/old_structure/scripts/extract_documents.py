#!/usr/bin/env python3
"""
Document Text Extractor - Offline Version
Extracts text from local documents (PDF, DOCX, etc.)
No internet required - works with files already downloaded.
"""

import os
import json
import logging
from pathlib import Path
from datetime import datetime
from dataclasses import dataclass, asdict
from typing import List, Dict, Optional, Any
from concurrent.futures import ThreadPoolExecutor, as_completed

# Document processing imports
import fitz  # PyMuPDF for PDFs
from docx import Document
from openpyxl import load_workbook
import csv

import logging

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# ============================================================
# CONFIGURATION - Edit these paths as needed
# ============================================================

# Where to look for documents (will search recursively)
INPUT_FOLDERS = [
    Path("data/downloads"),
    Path("data/downloads/processos"),
    Path("downloads"),
]

# Where to save extraction results
OUTPUT_DIR = Path("data/extractions")

# Supported file formats
SUPPORTED_EXTENSIONS = ['.pdf', '.docx', '.doc', '.txt', '.xlsx', '.xls', '.csv', '.md']

# Processing options
MAX_WORKERS = 4  # Parallel processing threads

# ============================================================

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.StreamHandler(),
        logging.FileHandler('extraction.log', encoding='utf-8')
    ]
)
logger = logging.getLogger(__name__)


@dataclass
class ExtractedDocument:
    """Represents an extracted document with content and metadata."""
    filename: str
    filepath: str
    file_type: str
    content: str
    metadata: Dict[str, Any]
    page_count: Optional[int] = None
    word_count: int = 0
    char_count: int = 0
    extracted_at: str = ""
    extraction_status: str = "success"
    error_message: str = ""
    
    def __post_init__(self):
        if not self.extracted_at:
            self.extracted_at = datetime.now().isoformat()
        self.word_count = len(self.content.split()) if self.content else 0
        self.char_count = len(self.content) if self.content else 0


class LocalDocumentExtractor:
    """Extracts text from local documents without internet."""
    
    def __init__(self):
        self.results: List[ExtractedDocument] = []
        OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    
    def find_documents(self) -> List[Path]:
        """Find all supported documents in input folders."""
        documents = []
        
        for folder in INPUT_FOLDERS:
            if not folder.exists():
                logger.warning(f"Folder not found: {folder}")
                continue
            
            # Search recursively
            for ext in SUPPORTED_EXTENSIONS:
                found = list(folder.rglob(f"*{ext}"))
                documents.extend(found)
                if found:
                    logger.info(f"Found {len(found)} {ext} files in {folder}")
        
        # Remove duplicates
        documents = list(set(documents))
        logger.info(f"Total documents found: {len(documents)}")
        
        return documents
    
    def extract_pdf(self, filepath: Path) -> ExtractedDocument:
        """Extract text from PDF file."""
        try:
            doc = fitz.open(str(filepath))
            
            text_content = []
            for page_num, page in enumerate(doc, 1):
                text = page.get_text("text")
                if text.strip():
                    text_content.append(f"--- Página {page_num} ---\n{text}")
            
            metadata = {
                "title": doc.metadata.get("title", "") or "",
                "author": doc.metadata.get("author", "") or "",
                "subject": doc.metadata.get("subject", "") or "",
                "creator": doc.metadata.get("creator", "") or "",
                "creation_date": doc.metadata.get("creationDate", "") or "",
                "modification_date": doc.metadata.get("modDate", "") or "",
            }
            
            page_count = len(doc)
            content = "\n\n".join(text_content)
            doc.close()
            
            return ExtractedDocument(
                filename=filepath.name,
                filepath=str(filepath.absolute()),
                file_type="pdf",
                content=content,
                metadata=metadata,
                page_count=page_count
            )
            
        except Exception as e:
            logger.error(f"Error extracting PDF {filepath.name}: {e}")
            return ExtractedDocument(
                filename=filepath.name,
                filepath=str(filepath.absolute()),
                file_type="pdf",
                content="",
                metadata={},
                extraction_status="error",
                error_message=str(e)
            )
    
    def extract_docx(self, filepath: Path) -> ExtractedDocument:
        """Extract text from Word document."""
        try:
            doc = Document(str(filepath))
            
            # Extract paragraphs
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            
            # Extract tables
            tables_text = []
            for table in doc.tables:
                rows = []
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    rows.append(" | ".join(cells))
                tables_text.append("\n".join(rows))
            
            content = "\n\n".join(paragraphs)
            if tables_text:
                content += "\n\n--- Tabelas ---\n" + "\n\n".join(tables_text)
            
            # Core properties
            props = doc.core_properties
            metadata = {
                "title": props.title or "",
                "author": props.author or "",
                "subject": props.subject or "",
                "created": str(props.created) if props.created else "",
                "modified": str(props.modified) if props.modified else "",
            }
            
            return ExtractedDocument(
                filename=filepath.name,
                filepath=str(filepath.absolute()),
                file_type="docx",
                content=content,
                metadata=metadata
            )
            
        except Exception as e:
            logger.error(f"Error extracting DOCX {filepath.name}: {e}")
            return ExtractedDocument(
                filename=filepath.name,
                filepath=str(filepath.absolute()),
                file_type="docx",
                content="",
                metadata={},
                extraction_status="error",
                error_message=str(e)
            )
    
    def extract_text_file(self, filepath: Path) -> ExtractedDocument:
        """Extract content from plain text files."""
        encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
        content = ""
        used_encoding = "unknown"
        
        for encoding in encodings:
            try:
                content = filepath.read_text(encoding=encoding)
                used_encoding = encoding
                break
            except UnicodeDecodeError:
                continue
        
        metadata = {
            "encoding": used_encoding,
            "size_bytes": filepath.stat().st_size,
        }
        
        return ExtractedDocument(
            filename=filepath.name,
            filepath=str(filepath.absolute()),
            file_type="text",
            content=content,
            metadata=metadata
        )
    
    def extract_excel(self, filepath: Path) -> ExtractedDocument:
        """Extract data from Excel files."""
        try:
            workbook = load_workbook(str(filepath), read_only=True, data_only=True)
            
            all_content = []
            sheet_info = []
            
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                rows = []
                
                for row in sheet.iter_rows(values_only=True):
                    if any(cell is not None for cell in row):
                        row_text = " | ".join(str(c) if c is not None else "" for c in row)
                        rows.append(row_text)
                
                if rows:
                    all_content.append(f"=== Planilha: {sheet_name} ===\n" + "\n".join(rows))
                    sheet_info.append({"name": sheet_name, "rows": len(rows)})
            
            workbook.close()
            
            metadata = {
                "sheets": sheet_info,
                "sheet_count": len(sheet_info)
            }
            
            return ExtractedDocument(
                filename=filepath.name,
                filepath=str(filepath.absolute()),
                file_type="excel",
                content="\n\n".join(all_content),
                metadata=metadata
            )
            
        except Exception as e:
            logger.error(f"Error extracting Excel {filepath.name}: {e}")
            return ExtractedDocument(
                filename=filepath.name,
                filepath=str(filepath.absolute()),
                file_type="excel",
                content="",
                metadata={},
                extraction_status="error",
                error_message=str(e)
            )
    
    def extract_csv_file(self, filepath: Path) -> ExtractedDocument:
        """Extract data from CSV files."""
        try:
            rows = []
            
            with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
                sample = f.read(4096)
                f.seek(0)
                
                try:
                    dialect = csv.Sniffer().sniff(sample, delimiters=',;\t|')
                except csv.Error:
                    dialect = csv.excel
                
                reader = csv.reader(f, dialect)
                for row in reader:
                    rows.append(" | ".join(row))
            
            metadata = {
                "row_count": len(rows),
                "delimiter": getattr(dialect, 'delimiter', ',')
            }
            
            return ExtractedDocument(
                filename=filepath.name,
                filepath=str(filepath.absolute()),
                file_type="csv",
                content="\n".join(rows),
                metadata=metadata
            )
            
        except Exception as e:
            logger.error(f"Error extracting CSV {filepath.name}: {e}")
            return ExtractedDocument(
                filename=filepath.name,
                filepath=str(filepath.absolute()),
                file_type="csv",
                content="",
                metadata={},
                extraction_status="error",
                error_message=str(e)
            )
    
    def extract_document(self, filepath: Path) -> Optional[ExtractedDocument]:
        """Extract content from a single document based on type."""
        suffix = filepath.suffix.lower()
        
        logger.info(f"Extracting: {filepath.name}")
        
        if suffix == '.pdf':
            return self.extract_pdf(filepath)
        elif suffix in ['.docx', '.doc']:
            return self.extract_docx(filepath)
        elif suffix in ['.txt', '.md']:
            return self.extract_text_file(filepath)
        elif suffix in ['.xlsx', '.xls']:
            return self.extract_excel(filepath)
        elif suffix == '.csv':
            return self.extract_csv_file(filepath)
        else:
            logger.warning(f"Unsupported format: {suffix}")
            return None
    
    def run(self):
        """Run the extraction pipeline."""
        logging.info("\n" + "=" * 60)
        logging.info("📄 EXTRATOR DE DOCUMENTOS LOCAIS")
        logging.info("=" * 60)
        logging.info(f"📂 Pastas de busca: {[str(f) for f in INPUT_FOLDERS]}")
        logging.info(f"📁 Saída: {OUTPUT_DIR}")
        logging.info("=" * 60 + "\n")
        
        # Find all documents
        documents = self.find_documents()
        
        if not documents:
            logger.warning("\n⚠️  Nenhum documento encontrado nas pastas configuradas!")
            logging.info("\n📝 Verifique se os documentos estão em:")
            for folder in INPUT_FOLDERS:
                logging.info(f"   - {folder.absolute()}")
            return
        
        # Process documents in parallel
        logging.info(f"\n🔄 Processando {len(documents)} documentos...\n")
        
        with ThreadPoolExecutor(max_workers=MAX_WORKERS) as executor:
            futures = {executor.submit(self.extract_document, doc): doc for doc in documents}
            
            completed = 0
            for future in as_completed(futures):
                result = future.result()
                if result:
                    self.results.append(result)
                completed += 1
                
                # Progress indicator
                if completed % 10 == 0 or completed == len(documents):
                    logging.info(f"   Progresso: {completed}/{len(documents)}")
        
        # Save results
        self.save_results()
        
        # Print summary
        self.print_summary()
    
    def save_results(self):
        """Save extraction results to JSON and CSV."""
        if not self.results:
            return
        
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        
        # Save detailed JSON
        json_file = OUTPUT_DIR / f"extracted_texts_{timestamp}.json"
        with open(json_file, 'w', encoding='utf-8') as f:
            json.dump(
                [asdict(r) for r in self.results],
                f,
                indent=2,
                ensure_ascii=False
            )
        
        # Save summary CSV
        csv_file = OUTPUT_DIR / f"extraction_summary_{timestamp}.csv"
        csv_data = []
        for r in self.results:
            csv_data.append({
                'Arquivo': r.filename,
                'Tipo': r.file_type,
                'Páginas': r.page_count or '-',
                'Palavras': r.word_count,
                'Caracteres': r.char_count,
                'Status': r.extraction_status,
                'Erro': r.error_message,
                'Caminho': r.filepath
            })
        
        import pandas as pd
        pd.DataFrame(csv_data).to_csv(csv_file, index=False, encoding='utf-8-sig')
        
        # Save combined text file (all content in one file)
        txt_file = OUTPUT_DIR / f"all_texts_{timestamp}.txt"
        with open(txt_file, 'w', encoding='utf-8') as f:
            for r in self.results:
                if r.content:
                    f.write("=" * 80 + "\n")
                    f.write(f"ARQUIVO: {r.filename}\n")
                    f.write(f"TIPO: {r.file_type}\n")
                    f.write(f"PALAVRAS: {r.word_count}\n")
                    f.write("=" * 80 + "\n\n")
                    f.write(r.content)
                    f.write("\n\n")
        
        logger.info(f"Saved: {json_file.name}")
        logger.info(f"Saved: {csv_file.name}")
        logger.info(f"Saved: {txt_file.name}")
    
    def print_summary(self):
        """Print extraction summary."""
        logging.info("\n" + "=" * 60)
        logging.info("📊 RESUMO DA EXTRAÇÃO")
        logging.info("=" * 60)
        
        total = len(self.results)
        success = sum(1 for r in self.results if r.extraction_status == "success")
        errors = sum(1 for r in self.results if r.extraction_status == "error")
        total_words = sum(r.word_count for r in self.results)
        total_pages = sum(r.page_count or 0 for r in self.results)
        
        # Count by type
        by_type = {}
        for r in self.results:
            by_type[r.file_type] = by_type.get(r.file_type, 0) + 1
        
        logging.info(f"\n📁 Total de arquivos: {total}")
        logging.info(f"✅ Extraídos com sucesso: {success}")
        logger.error(f"❌ Erros: {errors}")
        logging.info(f"📄 Total de páginas: {total_pages}")
        logging.info(f"📝 Total de palavras: {total_words:,}")
        
        logging.info(f"\n📈 Por tipo de arquivo:")
        type_icons = {
            'pdf': '📕',
            'docx': '📘',
            'excel': '📗',
            'csv': '📊',
            'text': '📄'
        }
        for file_type, count in sorted(by_type.items()):
            icon = type_icons.get(file_type, '📎')
            logging.info(f"   {icon} {file_type}: {count}")
        
        logging.info(f"\n📂 Resultados salvos em: {OUTPUT_DIR.absolute()}")
        logging.info("=" * 60)


def main():
    """Main entry point - no arguments needed."""
    extractor = LocalDocumentExtractor()
    extractor.run()


if __name__ == "__main__":
    main()